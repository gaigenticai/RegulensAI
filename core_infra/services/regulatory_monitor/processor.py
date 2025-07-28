"""
Document Processor - File Format Handler
Processes and extracts text from various document formats.
"""

import asyncio
import aiofiles
import aiohttp
from io import BytesIO
from typing import Optional, Dict, Any, List, Tuple
import re
from datetime import datetime
import structlog
import hashlib

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# HTML processing
try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

# Office document processing
try:
    from docx import Document as DocxDocument
    import openpyxl
    OFFICE_AVAILABLE = True
except ImportError:
    OFFICE_AVAILABLE = False

from core_infra.services.monitoring import track_operation
from core_infra.config import get_settings

logger = structlog.get_logger(__name__)
settings = get_settings()


class DocumentProcessor:
    """
    Processes various document formats and extracts clean text content
    for regulatory analysis.
    """
    
    def __init__(self):
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024  # Convert MB to bytes
        self.allowed_types = settings.allowed_file_types_list
        self.session = None
    
    async def __aenter__(self):
        """Async context manager entry."""
        timeout = aiohttp.ClientTimeout(total=60)
        self.session = aiohttp.ClientSession(timeout=timeout)
        return self
    
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """Async context manager exit."""
        if self.session:
            await self.session.close()
    
    @track_operation("document_processor.process_document")
    async def process_document(self, file_path: str = None, file_url: str = None, 
                             file_content: bytes = None, content_type: str = None) -> Dict[str, Any]:
        """
        Process a document from file path, URL, or content bytes.
        Returns extracted text and metadata.
        """
        try:
            # Determine processing method
            if file_content:
                return await self._process_content(file_content, content_type)
            elif file_url:
                return await self._process_url(file_url)
            elif file_path:
                return await self._process_file(file_path)
            else:
                raise ValueError("Must provide either file_path, file_url, or file_content")
                
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {
                'text': '',
                'metadata': {},
                'error': str(e),
                'processed_at': datetime.utcnow().isoformat()
            }
    
    async def _process_file(self, file_path: str) -> Dict[str, Any]:
        """Process document from local file path."""
        try:
            async with aiofiles.open(file_path, 'rb') as file:
                content = await file.read()
            
            # Determine content type from file extension
            content_type = self._get_content_type_from_path(file_path)
            
            return await self._process_content(content, content_type)
            
        except Exception as e:
            logger.error(f"Failed to process file {file_path}: {e}")
            raise
    
    async def _process_url(self, url: str) -> Dict[str, Any]:
        """Process document from URL."""
        try:
            if not self.session:
                async with self:
                    return await self._download_and_process(url)
            else:
                return await self._download_and_process(url)
                
        except Exception as e:
            logger.error(f"Failed to process URL {url}: {e}")
            raise
    
    async def _download_and_process(self, url: str) -> Dict[str, Any]:
        """Download document from URL and process."""
        try:
            headers = {
                'User-Agent': 'Regulens-AI-Document-Processor/1.0'
            }
            
            async with self.session.get(url, headers=headers) as response:
                if response.status != 200:
                    raise Exception(f"HTTP {response.status} when downloading {url}")
                
                # Check file size
                content_length = response.headers.get('content-length')
                if content_length and int(content_length) > self.max_file_size:
                    raise Exception(f"File too large: {content_length} bytes")
                
                content = await response.read()
                content_type = response.headers.get('content-type', '')
                
                return await self._process_content(content, content_type)
                
        except Exception as e:
            logger.error(f"Failed to download and process {url}: {e}")
            raise
    
    async def _process_content(self, content: bytes, content_type: str) -> Dict[str, Any]:
        """Process document content based on content type."""
        try:
            # Validate file size
            if len(content) > self.max_file_size:
                raise Exception(f"File too large: {len(content)} bytes")
            
            # Normalize content type
            content_type = content_type.lower().split(';')[0] if content_type else ''
            
            # Extract text based on content type
            if content_type in ['application/pdf'] or self._is_pdf_content(content):
                text, metadata = await self._extract_from_pdf(content)
            elif content_type in ['text/html', 'application/xhtml+xml']:
                text, metadata = await self._extract_from_html(content)
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                text, metadata = await self._extract_from_docx(content)
            elif content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
                text, metadata = await self._extract_from_xlsx(content)
            elif content_type in ['text/plain', 'text/csv']:
                text, metadata = await self._extract_from_text(content)
            elif content_type in ['application/json']:
                text, metadata = await self._extract_from_json(content)
            else:
                # Try to detect format from content
                text, metadata = await self._extract_auto_detect(content)
            
            # Clean and process extracted text
            cleaned_text = self._clean_text(text)
            
            # Generate document fingerprint
            fingerprint = self._generate_fingerprint(content)
            
            # Extract additional metadata
            extracted_metadata = self._extract_metadata(cleaned_text)
            metadata.update(extracted_metadata)
            
            return {
                'text': cleaned_text,
                'metadata': {
                    **metadata,
                    'content_type': content_type,
                    'file_size': len(content),
                    'fingerprint': fingerprint,
                    'word_count': len(cleaned_text.split()),
                    'char_count': len(cleaned_text),
                    'processed_at': datetime.utcnow().isoformat()
                },
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Content processing failed: {e}")
            return {
                'text': '',
                'metadata': {
                    'error': str(e),
                    'content_type': content_type,
                    'file_size': len(content) if content else 0,
                    'processed_at': datetime.utcnow().isoformat()
                },
                'success': False
            }
    
    async def _extract_from_pdf(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF content."""
        if not PDF_AVAILABLE:
            raise Exception("PDF processing libraries not available")
        
        try:
            text = ""
            metadata = {}
            
            # Try pdfplumber first (better text extraction)
            try:
                import pdfplumber
                with pdfplumber.open(BytesIO(content)) as pdf:
                    pages = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pages.append(page_text)
                    
                    text = '\n\n'.join(pages)
                    metadata.update({
                        'page_count': len(pdf.pages),
                        'extractor': 'pdfplumber'
                    })
                    
                    # Extract PDF metadata
                    if pdf.metadata:
                        metadata.update({
                            'title': pdf.metadata.get('Title', ''),
                            'author': pdf.metadata.get('Author', ''),
                            'creator': pdf.metadata.get('Creator', ''),
                            'creation_date': pdf.metadata.get('CreationDate', ''),
                            'modification_date': pdf.metadata.get('ModDate', '')
                        })
                        
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed, trying PyPDF2: {e}")
                
                # Fallback to PyPDF2
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                pages = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            pages.append(page_text)
                    except Exception as page_error:
                        logger.warning(f"Failed to extract page {page_num}: {page_error}")
                
                text = '\n\n'.join(pages)
                metadata.update({
                    'page_count': len(pdf_reader.pages),
                    'extractor': 'pypdf2'
                })
                
                # Extract PDF metadata
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'creation_date': pdf_reader.metadata.get('/CreationDate', ''),
                        'modification_date': pdf_reader.metadata.get('/ModDate', '')
                    })
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            return "", {"error": str(e)}
    
    async def _extract_from_html(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from HTML content."""
        if not HTML_AVAILABLE:
            # Fallback to simple regex
            return self._extract_html_regex(content)
        
        try:
            # Decode content
            html_content = content.decode('utf-8', errors='ignore')
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Extract metadata
            metadata = {}
            
            # Title
            title_tag = soup.find('title')
            if title_tag:
                metadata['title'] = title_tag.get_text().strip()
            
            # Meta tags
            meta_description = soup.find('meta', attrs={'name': 'description'})
            if meta_description:
                metadata['description'] = meta_description.get('content', '')
            
            meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
            if meta_keywords:
                metadata['keywords'] = meta_keywords.get('content', '')
            
            # Extract headings structure
            headings = []
            for i in range(1, 7):
                for heading in soup.find_all(f'h{i}'):
                    headings.append({
                        'level': i,
                        'text': heading.get_text().strip()
                    })
            
            if headings:
                metadata['headings'] = headings[:20]  # Limit to first 20 headings
            
            # Extract links
            links = []
            for link in soup.find_all('a', href=True):
                link_text = link.get_text().strip()
                if link_text and len(link_text) > 5:  # Filter out short/empty links
                    links.append({
                        'text': link_text,
                        'href': link['href']
                    })
            
            if links:
                metadata['links'] = links[:50]  # Limit to first 50 links
            
            metadata['extractor'] = 'beautifulsoup'
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            return self._extract_html_regex(content)
    
    def _extract_html_regex(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Fallback HTML text extraction using regex."""
        try:
            html_content = content.decode('utf-8', errors='ignore')
            
            # Remove script and style tags with content
            html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', '', html_content)
            
            # Extract title
            title_match = re.search(r'<title[^>]*>(.*?)</title>', html_content, re.IGNORECASE | re.DOTALL)
            title = title_match.group(1).strip() if title_match else ''
            
            metadata = {
                'title': title,
                'extractor': 'regex'
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Regex HTML extraction failed: {e}")
            return "", {"error": str(e)}
    
    async def _extract_from_docx(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX content."""
        if not OFFICE_AVAILABLE:
            raise Exception("Office document processing libraries not available")
        
        try:
            doc = DocxDocument(BytesIO(content))
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(' | '.join(row_text))
            
            # Combine all text
            all_text = paragraphs + table_text
            text = '\n\n'.join(all_text)
            
            # Extract metadata
            metadata = {
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'extractor': 'python-docx'
            }
            
            # Document properties
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': core_props.created.isoformat() if core_props.created else '',
                    'modified': core_props.modified.isoformat() if core_props.modified else ''
                })
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return "", {"error": str(e)}
    
    async def _extract_from_xlsx(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from XLSX content."""
        if not OFFICE_AVAILABLE:
            raise Exception("Office document processing libraries not available")
        
        try:
            workbook = openpyxl.load_workbook(BytesIO(content), data_only=True)
            
            all_text = []
            sheet_info = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                row_count = 0
                
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = []
                        for cell in row:
                            if cell is not None:
                                row_text.append(str(cell))
                        if row_text:
                            sheet_text.append(' | '.join(row_text))
                            row_count += 1
                
                if sheet_text:
                    all_text.append(f"=== Sheet: {sheet_name} ===\n" + '\n'.join(sheet_text))
                    sheet_info.append({
                        'name': sheet_name,
                        'rows': row_count
                    })
            
            text = '\n\n'.join(all_text)
            
            metadata = {
                'sheet_count': len(workbook.sheetnames),
                'sheets': sheet_info,
                'extractor': 'openpyxl'
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"XLSX extraction failed: {e}")
            return "", {"error": str(e)}
    
    async def _extract_from_text(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text or CSV content."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    metadata = {
                        'encoding': encoding,
                        'extractor': 'text'
                    }
                    return text, metadata
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use errors='ignore'
            text = content.decode('utf-8', errors='ignore')
            metadata = {
                'encoding': 'utf-8-ignore',
                'extractor': 'text'
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return "", {"error": str(e)}
    
    async def _extract_from_json(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from JSON content."""
        try:
            import json
            
            json_content = content.decode('utf-8')
            data = json.loads(json_content)
            
            # Convert JSON to readable text
            text = self._json_to_text(data)
            
            metadata = {
                'extractor': 'json',
                'json_structure': self._analyze_json_structure(data)
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"JSON extraction failed: {e}")
            return "", {"error": str(e)}
    
    def _json_to_text(self, data, prefix="") -> str:
        """Convert JSON data to readable text."""
        lines = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{prefix}{key}:")
                    lines.append(self._json_to_text(value, prefix + "  "))
                else:
                    lines.append(f"{prefix}{key}: {value}")
        elif isinstance(data, list):
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{prefix}[{i}]:")
                    lines.append(self._json_to_text(item, prefix + "  "))
                else:
                    lines.append(f"{prefix}[{i}]: {item}")
        else:
            lines.append(f"{prefix}{data}")
        
        return '\n'.join(lines)
    
    def _analyze_json_structure(self, data) -> Dict[str, Any]:
        """Analyze JSON structure for metadata."""
        if isinstance(data, dict):
            return {
                'type': 'object',
                'keys': list(data.keys())[:20],  # Limit to first 20 keys
                'key_count': len(data)
            }
        elif isinstance(data, list):
            return {
                'type': 'array',
                'length': len(data),
                'item_types': list(set(type(item).__name__ for item in data[:10]))  # Sample first 10 items
            }
        else:
            return {
                'type': type(data).__name__,
                'value': str(data)[:100]  # First 100 chars
            }
    
    async def _extract_auto_detect(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Auto-detect format and extract text."""
        try:
            # Check if it's PDF
            if self._is_pdf_content(content):
                return await self._extract_from_pdf(content)
            
            # Check if it's HTML
            if self._is_html_content(content):
                return await self._extract_from_html(content)
            
            # Default to text extraction
            return await self._extract_from_text(content)
            
        except Exception as e:
            logger.error(f"Auto-detection failed: {e}")
            return "", {"error": str(e)}
    
    def _is_pdf_content(self, content: bytes) -> bool:
        """Check if content is PDF format."""
        return content.startswith(b'%PDF-')
    
    def _is_html_content(self, content: bytes) -> bool:
        """Check if content is HTML format."""
        try:
            text = content[:1000].decode('utf-8', errors='ignore').lower()
            return any(tag in text for tag in ['<html', '<head', '<body', '<!doctype html'])
        except:
            return False
    
    def _get_content_type_from_path(self, file_path: str) -> str:
        """Get content type from file extension."""
        extension = file_path.lower().split('.')[-1] if '.' in file_path else ''
        
        content_type_map = {
            'pdf': 'application/pdf',
            'html': 'text/html',
            'htm': 'text/html',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'txt': 'text/plain',
            'csv': 'text/csv',
            'json': 'application/json'
        }
        
        return content_type_map.get(extension, 'application/octet-stream')
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'Page \d+ of \d+', '', text)
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Clean up special characters
        text = text.replace('\x00', '')  # Remove null characters
        text = text.replace('\r\n', '\n')  # Normalize line endings
        text = text.replace('\r', '\n')
        
        # Remove excessive spaces
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Trim
        text = text.strip()
        
        return text
    
    def _generate_fingerprint(self, content: bytes) -> str:
        """Generate unique fingerprint for document content."""
        return hashlib.sha256(content).hexdigest()
    
    def _extract_metadata(self, text: str) -> Dict[str, Any]:
        """Extract additional metadata from text content."""
        metadata = {}
        
        # Extract dates
        date_patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{4}\b',  # MM/DD/YYYY or MM-DD-YYYY
            r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',  # YYYY/MM/DD or YYYY-MM-DD
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            dates.extend(matches)
        
        if dates:
            metadata['extracted_dates'] = list(set(dates))[:10]  # Limit to 10 unique dates
        
        # Extract email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            metadata['extracted_emails'] = list(set(emails))[:5]  # Limit to 5 unique emails
        
        # Extract phone numbers (simple pattern)
        phone_pattern = r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b'
        phones = re.findall(phone_pattern, text)
        if phones:
            metadata['extracted_phones'] = list(set(phones))[:5]  # Limit to 5 unique phones
        
        # Extract regulatory references (e.g., "Section 123", "Rule 456")
        reg_pattern = r'\b(?:Section|Rule|Part|Article|Regulation)\s+\d+[A-Za-z]?\b'
        regulations = re.findall(reg_pattern, text, re.IGNORECASE)
        if regulations:
            metadata['regulatory_references'] = list(set(regulations))[:10]
        
        return metadata
    
    async def extract_text_from_html(self, html_content: str) -> str:
        """Public method to extract text from HTML string."""
        content_bytes = html_content.encode('utf-8')
        text, _ = await self._extract_from_html(content_bytes)
        return self._clean_text(text)
    
    async def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """Public method to extract text from PDF bytes."""
        text, _ = await self._extract_from_pdf(pdf_content)
        return self._clean_text(text) 